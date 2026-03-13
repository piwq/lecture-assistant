from flask import Flask, request, jsonify, send_from_directory, Response
from google import genai
import docx
import os
import json
import concurrent.futures
import threading

app = Flask(__name__)

# --- CONFIG ---
DOCX_FILENAME = "lecture.docx"
doc_lock = threading.Lock() # Блокировка для безопасной записи из разных потоков

# Полный список доступных текстовых моделей из твоего API
MODELS_TO_QUERY = [
    # Флагманы и стабильные версии
    'gemini-2.5-flash',
    'gemini-2.5-pro',
    
    # Алиасы на самые свежие стабильные версии
    'gemini-flash-latest',
    'gemini-pro-latest',
    
    # Превью новых 3-х поколений (самые умные)
    'gemini-3-flash-preview',
    'gemini-3.1-pro-preview',
    'gemini-3.1-flash-lite-preview',
    
    # Локальные/Открытые модели (отвечают по-разному, интересно сравнить)
    'gemma-3-4b-it',
    'gemma-3-27b-it'
]
# --- HELPERS ---

def log_request_to_docx(lecture, question):
    """Записывает начало блока (Лекция + Вопрос) в файл перед ответами моделей."""
    with doc_lock:
        try:
            doc = docx.Document(DOCX_FILENAME) if os.path.exists(DOCX_FILENAME) else docx.Document()
            doc.add_heading('=' * 20 + ' НОВЫЙ ЗАПРОС ' + '=' * 20, level=1)
            doc.add_heading('Конспект лекции:', level=2)
            doc.add_paragraph(lecture if lecture else "(Текст лекции пуст)")
            doc.add_heading('Вопрос преподавателя:', level=2)
            doc.add_paragraph(question if question else "(Вопрос не распознан)")
            doc.add_heading('Ответы нейросетей:', level=2)
            doc.save(DOCX_FILENAME)
        except Exception as e:
            print(f"❌ Ошибка записи лекции в DOCX: {e}")

def log_answer_to_docx(model_name, answer_text):
    """Аккуратно дописывает ответ конкретной модели в конец файла."""
    with doc_lock:
        try:
            doc = docx.Document(DOCX_FILENAME) if os.path.exists(DOCX_FILENAME) else docx.Document()
            # Добавляем жирным шрифтом название модели и затем её ответ
            p = doc.add_paragraph()
            p.add_run(f"[{model_name}]: ").bold = True
            p.add_run(answer_text)
            doc.save(DOCX_FILENAME)
        except Exception as e:
            print(f"❌ Ошибка записи ответа в DOCX: {e}")
            
def generate_and_format(model_name, lecture, question, api_key):
    """Synchronously generates and formats an answer from a single model."""
    try:
        client = genai.Client(api_key=api_key)

        # 1. Жесткий русскоязычный промпт для генерации ответа
        prompt1 = f"""Контекст: Лекция. Студент должен ответить на вопрос преподавателя.
ТЕКСТ ЛЕКЦИИ:
{lecture}

ВОПРОС ПРЕПОДАВАТЕЛЯ:
{question}

ЗАДАЧА:
Дай КРАТКИЙ ответ (максимум 1-2 предложения), опираясь только на текст лекции.
Отвечай так, как будто студент произносит это вслух.
ВНИМАНИЕ: Не используй вводные слова. Не пиши "Ответ:". Напиши только сам факт."""
        
        response1 = client.models.generate_content(
            model=model_name,
            contents=prompt1,
        )
        raw_answer = response1.text

        # 2. Жесткий русскоязычный промпт для форматирования
        prompt2 = f"""Исправь пунктуацию и орфографию в тексте ниже. Сделай его читаемым.
ВНИМАНИЕ: Выведи ТОЛЬКО исправленный текст. Никаких комментариев, никаких фраз вроде "Вот ваш текст" или перевода на английский.

ОРИГИНАЛЬНЫЙ ТЕКСТ:
{raw_answer}"""
        
        response2 = client.models.generate_content(
            model=model_name,
            contents=prompt2,
        )
        formatted_answer = response2.text.strip()
        
        log_answer_to_docx(model_name, formatted_answer)

        return {"model": model_name, "answer": formatted_answer}

    except Exception as e:
        print(f"❌ Ошибка от модели {model_name}: {e}")
        return {"model": model_name, "error": str(e)}

# --- ROUTES ---

@app.route('/')
def serve_index():
    return send_from_directory('.', 'lecture.html')

@app.route('/ask', methods=['POST'])
def ask():
    data = request.get_json()
    api_key = data.get('apiKey')
    lecture = data.get('lecture')
    question = data.get('question')

    if not all([api_key, lecture, question]):
        return "Missing required fields", 400

    log_request_to_docx(lecture, question)
    
    def event_stream():
        with concurrent.futures.ThreadPoolExecutor(max_workers=len(MODELS_TO_QUERY)) as executor:
            futures = [
                executor.submit(generate_and_format, model, lecture, question, api_key) 
                for model in MODELS_TO_QUERY
            ]
            
            try:
                for future in concurrent.futures.as_completed(futures):
                    result = future.result()
                    yield f"data: {json.dumps(result)}\n\n"
            except Exception as e:
                print(f"❌ Ошибка в стриме: {e}")
                error_payload = json.dumps({'error': str(e), 'model': 'Stream Error'})
                yield f"data: {error_payload}\n\n"

    return Response(event_stream(), mimetype='text/event-stream')

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=3333, debug=True, threaded=True)